from datetime import datetime
from flask import Flask, render_template, url_for, request, redirect, flash, session
from flask_session import Session
from flask_sqlalchemy import SQLAlchemy
from flask_login import UserMixin, login_user, LoginManager, login_required, current_user, logout_user
from flask_wtf import FlaskForm, CSRFProtect
from flask_wtf.csrf import CSRFError
from wtforms import BooleanField, StringField, PasswordField, SubmitField, validators
from wtforms.validators import InputRequired, DataRequired, Email, Length, EqualTo, ValidationError
from flask_bcrypt import Bcrypt
from dotenv import load_dotenv
from langchain_text_splitters import CharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from langchain_community.chains import RetrievalQA
import os
import fitz  # PyMuPDF
import docx
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain_xai import ChatXAI
from langchain.chains import ConversationalRetrievalChain
from langchain.prompts import PromptTemplate
from flask_bootstrap5 import Bootstrap
from dotenv import load_dotenv
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask import jsonify


load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '.env'))



load_dotenv()

xai_key = os.getenv("XAI_API_KEY")
#os.environ["XAI_API_KEY"] = xai_key

app = Flask(__name__, static_url_path='/static', static_folder='static')


# Mail config
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'true').lower() in ['true', '1']
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
mail = Mail(app)



limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["20 per minute", "100 per hour"]
)

app.config['SECRET_KEY'] = 'secret_key'

# Token serializer
s = URLSafeTimedSerializer(app.config['SECRET_KEY'])

def generate_reset_token(email):
    return s.dumps(email, salt='password-reset')

def verify_reset_token(token, expiration=3600):
    try:
        return s.loads(token, salt='password-reset', max_age=expiration)
    except Exception:
        return None

#app = Flask(__name__)

# Connecting to the database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
# Adding secret key for session management
#app.config['SECRET_KEY'] = 'secret_key'

app.config['SESSION_TYPE'] = 'filesystem'
#Session(app)

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
csrf = CSRFProtect(app)
Session(app)
Bootstrap(app)

@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    flash('CSRF token missing or invalid. Please try again.', 'danger')
    return redirect(url_for('home'))

# Initialize the login manager
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'  # Redirect to login page if not logged in

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))
    #return User.query.get(int(user_id))

# Initialize the SQLAlchemy object AFTER app config
#db = SQLAlchemy(app)
#bcrypt = Bcrypt(app)


# Creating user login table database
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True) # Identity Columu
    username = db.Column(db.String(150), unique=True, nullable=False) # Username Column
    email = db.Column(db.String(150), unique=True, nullable=False) # Email Column
    password = db.Column(db.String(150), nullable=False) # Password Column

class RegisterForm(FlaskForm):
    username = StringField(validators=[InputRequired(), Length(min=4, max=20)], render_kw={"placeholder": "Username"})
    email = StringField(validators=[InputRequired(), Email()], render_kw={"placeholder": "Email"})
    password = PasswordField(validators=[InputRequired(), Length(min=4, max=80)], render_kw={"placeholder": "Password"})
    confirm_password = PasswordField(validators=[InputRequired(), EqualTo('password', message='Passwords must match')], render_kw={"placeholder": "Confirm Password"})
    submit = SubmitField('Register')

    def validate_username(self, username):
        existing_user = User.query.filter_by(username=username.data).first()
        if existing_user:
            raise ValidationError('Username already exists. Please choose a different one.')
        
class LoginForm(FlaskForm):
    username = StringField(validators=[InputRequired(), Length(min=4, max=20)], render_kw={"placeholder": "Username"})
    #email = StringField(validators=[InputRequired(), Email()], render_kw={"placeholder": "Email"})
    password = PasswordField(validators=[InputRequired(), Length(min=4, max=80)], render_kw={"placeholder": "Password"})
    #confirm_password = PasswordField(validators=[InputRequired(), EqualTo('password', message='Passwords must match')], render_kw={"placeholder": "Confirm Password"})
    remember = BooleanField('Remember me')
    submit = SubmitField('Login')

class RequestResetForm(FlaskForm):
    email = StringField('Email', validators=[InputRequired(), Email()])
    submit = SubmitField('Request Password Reset')

class ResetPasswordForm(FlaskForm):
    password = PasswordField('New Password', validators=[InputRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[EqualTo('password')])
    submit = SubmitField('Reset Password')

class PasswordResetLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(150), nullable=False)
    ip_address = db.Column(db.String(45))
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    success = db.Column(db.Boolean, default=False)




xai_api_key = os.getenv('XAI_API_KEY')
print("XAI KEY:", xai_api_key)


# Home page — file upload form
#@app.route('/', methods=['GET', 'POST'])
#def home():
#    text = ""
#    answer = ""
#    question = ""

#    if request.method == 'POST':
#        uploaded_file = request.files.get('document')
#        question = request.form.get('question', '').strip()

#        if uploaded_file and uploaded_file.filename:
            # Extract text from file
#            if uploaded_file.filename.endswith(".pdf"):
#                doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
#                text = "".join([page.get_text() for page in doc])
#            elif uploaded_file.filename.endswith(".docx") or uploaded_file.filename.endswith(".doc"):
#                doc = docx.Document(uploaded_file)
#                text = "\n".join([para.text for para in doc.paragraphs])
#            elif uploaded_file.filename.endswith(".txt"):
#                text = uploaded_file.read().decode('utf-8')
#            else:
#                return render_template('result.html', extracted_text="", question=question, answer="Unsupported file type.")


@app.route('/', methods=['GET', 'POST'])
def home():
    text = session.get('document_text', '') # Store and retrieve text from session
    #uploaded_file = request.files['document']
    question = request.form.get('question', '')
         
    if request.method == 'POST': #and 'document' in request.files:
        # Check if a file is uploaded
        if 'document' in request.files and request.files['document'].filename:
            uploaded_file = request.files['document']
            print(f"Uploaded file: {uploaded_file.filename}")
            if uploaded_file:
                if uploaded_file.filename.endswith(".pdf"):
                    try:
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        text = ""
                        for page in doc:
                            text += page.get_text()
                        print(f"PDF text length: {len(text)}")
                    except Exception as e:
                        text = f"Error reading PDF file: {str(e)}"
                        print(f"PDF error: {text}")
                elif uploaded_file.filename.endswith((".docx", ".doc")):
                    doc = docx.Document(uploaded_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    print(f"DOC/DOCX text length: {len(text)}")
                #elif uploaded_file.filename.endswith(".doc"):
                #    doc = docx.Document(uploaded_file)
                #    print(f"DOC/DOCX text length: {len(text)}")
                #    text = "\n".join([para.text for para in doc.paragraphs])
                elif uploaded_file.filename.endswith(".txt"):
                    text = uploaded_file.read().decode('utf-8')  # Read text file content
                    print(f"TXT text length: {len(text)}")
                else:
                    text = "Unsupported file type"
                    print(f"Unsupported file type: {text}")
                # Truncate text to fit cookie size (e.g., 4000 characters) and update session
                #max_length = 9000  # Adjust based on testing
                #text = text[:max_length] if len(text) > max_length else text
                # Clear and update session with new document text
                #session.clear()  # Clear existing session data
                # Store the extracted text in the session
                
                if 'document_text' in session:
                    del session['document_text']  # Clear previous text if exists
                print(f"Storing text in session, length: {len(text)}")  #
                session['document_text'] = text
                print(f"Updated session text length: {len(session.get('document_text', ''))}")  # Debug: Session update
        #question = ""
        
            
        # Process the document and question as needed
        if question and text:
            # Step 1: Split into chunks
            splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            docs = [Document(page_content=chunk) for chunk in splitter.split_text(text)]
            print(f"Number of document chunks: {len(docs)}")  # Debug: Chunk count

            # Step 2: Embed chunks
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-MiniLM-L6-v2")
            db = FAISS.from_documents(docs, embeddings)
            print(f"Vector store recreated with {len(docs)} documents")  # Debug: Vector store creation

            # Step 3: Connect to XAI LLM
            # If LangChain supports XAI via ChatOpenAI-compatible wrapper:

            print("XAI final key:", xai_key)
            print("From ENV:", os.environ.get("XAI_API_KEY"))

            llm = ChatXAI(
                temperature=0.3,
                #xai_api_key=os.getenv("XAI_API_KEY"),
                #api_key=os.getenv("XAI_API_KEY"),  # Ensure you have set this environment variable
                api_key= xai_api_key,
                openai_api_base="https://api.x.ai/v1",   # Replace with actual XAI base URL
                model="grok-3-mini-fast"  # Replace with your actual model name
                )

            # Step 4: Retrieval-based QA
            qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=db.as_retriever()
            )

            # Step 5: Ask a question
            print("Processing question...")
            print(f"Current text length: {len(text)}")
            print(f"QUESTION: {question}")
            docs_preview = db.similarity_search(question, k=2)
            print("RETRIEVED DOCS:")
            for doc in docs_preview:
                print(doc.page_content[:500])

            try:
                answer = qa_chain.run(question) #replace run with invoke
            except Exception as e:
                answer = f"Error during QA: {str(e)}"

            # This was added
            #return render_template('result.html', extracted_text=text[:1000], question=question, answer=answer)
        #else:
        #    return render_template('result.html', extracted_text=text[:1000], question=question, answer="Please upload a document and enter a question.")
        
    #else:
        #return render_template('result.html', extracted_text="", question=question, answer="Please upload a document to proceed.")

        #return render_template('iii.html')
            
            #answer = ""
    
        #user_question = request.form.get('question')
        #if user_question:
            #answer = qa_chain.run(user_question)

            #answer = ""
            #if question:
                #answer = qa_chain.run(question)
            return render_template(
                'result.html',
                extracted_text=text[:1000],
                answer=answer,
                question= question
                )
    return render_template('iii.html')#, extracted_text=text[:1000] if text else '')#, question=question)

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and bcrypt.check_password_hash(user.password, form.password.data):
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'danger')
    else:
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"Error in {field}: {error}", 'danger')
    return render_template('login.html', form=form)

@app.route('/reset_password', methods=['GET', 'POST'])
def request_password_reset():
    form = RequestResetForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        reset_log = PasswordResetLog(
            email=form.email.data,
            ip_address=request.remote_addr,
            success=bool(user)  # True if user exists, False otherwise
        )
        db.session.add(reset_log)
        db.session.commit()
        
        if user:
            token = generate_reset_token(user.email)
            reset_url = url_for('reset_token', token=token, _external=True)
            msg = Message('Password Reset Request', recipients=[user.email], sender=app.config['MAIL_USERNAME'])
            msg.body = f'''To reset your password, visit the following link:
{reset_url}
If you didn't make this request, ignore this email.'''
            mail.send(msg)
        flash('If your email exists, a reset link has been sent.', 'info')
        return redirect(url_for('login'))
    return render_template('request_reset.html', form=form)

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_token(token):
    email = verify_reset_token(token)
    if not email:
        flash('That reset link is invalid or expired.', 'warning')
        return redirect(url_for('reset_request'))
    form = ResetPasswordForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=email).first()
        if user:
            hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
            user.password = hashed_password
            db.session.commit()
            flash('Your password has been updated!', 'success')
            return redirect(url_for('login'))
    return render_template('reset_token.html', form=form)



@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
def dashboard():
    if not current_user.is_authenticated:
        return redirect(url_for('login'))
    
    # Here you can add logic to display user-specific data
    return render_template('dashboard.html', username=current_user.username)

@app.route('/iii', methods=['GET', 'POST'])
@login_required
def iii():
    return render_template('iii.html')
    

@app.route('/logout', methods=['GET', 'POST'])
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegisterForm()

    if form.validate_on_submit():
        hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
        new_user = User(
            username=form.username.data,
            email=form.email.data,
            password=hashed_password
        )
        db.session.add(new_user)
        db.session.commit()
        flash('Registration successful! You can now log in.', 'success')
        # Automatically log in the user after registration
        return redirect(url_for('login'))
    else:
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"Error in {field}: {error}", 'danger')
    return render_template('register.html', form=form)


"""def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        # Here you would typically check the credentials against a database
        if username == 'admin' and password == 'password':  # Example check
            return "Login successful!"
        else:
            return "Invalid credentials"
    return render_template('login.html')"""

# ✅ Add this here, before app.run()
@app.errorhandler(429)
def ratelimit_handler(e):
    retry_after = e.description or "a few seconds"
    return render_template('429.html', message=f"Please try again after {retry_after}."), 429

# Initialize the database
with app.app_context():
    db.create_all()

#if __name__ == '__main__':
    #app.run(debug=True, use_reloader=False)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True, use_reloader=False)  # Set use_reloader=False to avoid double initialization
