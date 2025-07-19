import streamlit as st
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI 
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.schema import Document
import os
import fitz  # PyMuPDF
import docx
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_xai import ChatXAI
from langchain.chains import ConversationalRetrievalChain
from langchain.prompts import PromptTemplate


# Upload the document
uploaded_file = st.file_uploader("📄 Upload a PDF or DOCX file", type=["pdf", "docx"])

if uploaded_file:
    st.success("✅ File uploaded successfully!")

    if uploaded_file.name.endswith(".pdf"):
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
    elif uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        st.warning("Unsupported file type")
        text = ""

    st.text_area("📄 Extracted Text", text[:1000])

    # Ask for XAI API Key
    xai_api_key = st.text_input("xai_api_key")

    if xai_api_key:
        os.environ["xai_api_key"] = xai_api_key  # Optional, if the SDK uses env variable

        # Step 1: Split into chunks
        splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        docs = [Document(page_content=chunk) for chunk in splitter.split_text(text)]

        # Step 2: Embed chunks
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-MiniLM-L6-v2")
        db = FAISS.from_documents(docs, embeddings)

        # Step 3: Connect to XAI LLM
        # If LangChain supports XAI via ChatOpenAI-compatible wrapper:
        llm = ChatXAI(
            temperature=0.3,
            api_key='xai_api_key',
            openai_api_base="https://api.x.ai/v1",   # Replace with actual XAI base URL
            model="grok-3-mini-fast"  # Replace with your actual model name
        )

        # Step 4: Retrieval-based QA
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=db.as_retriever()
        )

        # Step 5: Ask a question
        user_question = st.text_input("💬 Ask a question about your course:")
        if user_question:
            answer = qa_chain.run(user_question)
            st.markdown("🎓 *Answer:* " + answer)

import streamlit as st

# Custom CSS
st.markdown("""
    <style>
        .main {
            background-color: #f0f4f8;
            background-image: url("https://images.unsplash.com/photo-1524995997946-a1c2e315a42f");
            background-size: cover;
            background-position: center;
            padding: 2rem;
            color: #fff;
        }

        h1, h2, h3 {
            color: #ffffff !important;
        }

        .stTextInput > div > div > input {
            background-color: #ffffff;
            color: #000;
        }

        .stTextArea > div > textarea {
            background-color: #ffffff;
            color: #000;
        }

        .css-1v0mbdj {
            background-color: rgba(255, 255, 255, 0.8);
            padding: 20px;
            border-radius: 15px;
        }
    </style>
""", unsafe_allow_html=True)

st.title("🎓 Your Study Companion")

col1, col2 = st.columns([1, 2])
with col1:
    st.image("https://yourlogo.url/logo.png", width=150)
with col2:
    st.title("Your Study Companion")

